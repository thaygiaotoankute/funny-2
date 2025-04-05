import os
import re
import json
import base64
import hashlib
import xml.etree.ElementTree as ET
import requests
import shutil
import time
import logging
from flask import Flask, request, jsonify, render_template, send_file
from werkzeug.utils import secure_filename
from Crypto.PublicKey import RSA
from Crypto.Cipher import PKCS1_v1_5
from PyPDF2 import PdfReader
from mistralai import Mistral

app = Flask(__name__)

# Vercel dùng hệ thống tệp tạm thời, thay đổi thư mục lưu trữ
if os.environ.get('VERCEL_ENV') == 'production':
    # Sử dụng thư mục /tmp trên Vercel
    app.config['UPLOAD_FOLDER'] = '/tmp'
else:
    app.config['UPLOAD_FOLDER'] = 'uploads'
    
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max upload size
# Tăng thời gian timeout cho các request
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 300  # 5 phút
app.config['PERMANENT_SESSION_LIFETIME'] = 1800  # 30 phút

# Tạo thư mục upload nếu chưa tồn tại
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Cấu hình logging
if os.environ.get('VERCEL_ENV') == 'production':
    app.logger.setLevel(logging.INFO)
else:
    app.logger.setLevel(logging.DEBUG)
    
# Log handler
handler = logging.StreamHandler()
handler.setFormatter(logging.Formatter(
    '%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
))
app.logger.addHandler(handler)

# Tăng thời gian timeout cho các request
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 300  # 5 phút

def call_gemini_api(original_text: str, gemini_key: str) -> str:
    """
    Gọi Gemini API để hiệu đính lỗi chính tả và ngữ pháp tiếng Việt.
    IMPORTANT: Không thay đổi image paths, công thức LaTeX và giữ nguyên dấu tiếng Việt.
    """
    try:
        if not gemini_key:
            return "Lỗi: Chưa có Gemini API Key"
        GEMINI_API_URL = (
            "https://generativelanguage.googleapis.com/v1beta/models/"
            "gemini-2.0-flash:generateContent?key=" + gemini_key
        )
        prompt = (
            "Please help me correct Vietnamese spelling and grammar in the following text. "
            "IMPORTANT: Do not change any image paths, LaTeX formulas, or Vietnamese diacritical marks. "
            "Return only the corrected text with the same structure and markdown formatting:\n\n"
            f"{original_text}"
        )
        payload = {
            "contents": [{
                "parts": [{
                    "text": prompt
                }]
            }],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 8192,
            }
        }
        headers = {"Content-Type": "application/json"}
        resp = requests.post(GEMINI_API_URL, json=payload, headers=headers, timeout=(30, 300))
        if resp.status_code == 200:
            data = resp.json()
            if "candidates" in data and len(data["candidates"]) > 0:
                candidate = data["candidates"][0]
                if "content" in candidate and "parts" in candidate["content"]:
                    corrected_text = candidate["content"]["parts"][0].get("text", "")
                    if corrected_text.strip():
                        return corrected_text
            return "Lỗi: Không thể trích xuất được kết quả từ Gemini API."
        else:
            return f"Lỗi: Gemini API - HTTP {resp.status_code} - {resp.text}"
    except Exception as e:
        return f"Lỗi: Gọi Gemini API thất bại: {str(e)}"

# Hàm tiện ích từ ứng dụng gốc (giữ nguyên các hàm này)
def load_rsa_private_key_from_xml(xml_str):
    """Tải khóa RSA riêng tư từ định dạng XML"""
    root = ET.fromstring(xml_str)
    def get_int(tag):
        text = root.find(tag).text
        return int.from_bytes(base64.b64decode(text), 'big')
    n = get_int('Modulus')
    e = get_int('Exponent')
    d = get_int('D')
    p = get_int('P')
    q = get_int('Q')
    key = RSA.construct((n, e, d, p, q))
    return key

def decrypt_api_key(encrypted_key_base64, rsa_private_key):
    """Giải mã API key đã được mã hóa"""
    try:
        cipher = PKCS1_v1_5.new(rsa_private_key)
        encrypted_data = base64.b64decode(encrypted_key_base64)
        decrypted = cipher.decrypt(encrypted_data, None)
        
        if not decrypted:
            raise ValueError("Giải mã thất bại")
        return decrypted.decode('utf-8')
    except Exception as e:
        raise ValueError(f"Lỗi giải mã API key: {str(e)}")

def get_mineru_token():
    """Lấy API key từ GitHub"""
    PRIVATE_KEY_XML = """<RSAKeyValue>
<Modulus>pWVItQwZ7NCPcBhSL4rqJrwh4OQquiPVtqTe4cqxO7o+UjYNzDPfLkfKAvR8k9ED4lq2TU11zEj8p2QZAM7obUlK4/HVexzfZd0qsXlCy5iaWoTQLXbVdzjvkC4mkO5TaX3Mpg/+p4oZjk1iS68tQFmju5cT19dcsPh554ICk8U=</Modulus>
<Exponent>AQAB</Exponent>
<P>0ZWwsKa9Vw9BJAsRaW4eV60i6Z+R6z9LNSgjNn4pYH2meZtGUbmJVowRv7EM5sytouB5EMru7sQbRHEQ7nrwSw==</P>
<Q>ygZQWNkUgfHhHBataXvYLxWgPB5UZTWogN8Mb33LT4rq7I5P1GX3oWtYF2AdmChX8Lq3Ms/A/jBhqYomhYOiLw==</Q>
<DP>qS9VOsTfA3Bk/VuR6rHh/JTfIgiWGnk1lOuZwVuGu0WzJWebFE3Z9+uKSFv8NjPz1w+tq0imKEhWWqGLMXg8kQ==</DP>
<DQ>UCtXQRrMB5EL6tCY+k4aCP1E+/ZxOUSk3Jcm4SuDPcp71WnYBgp8zULCz2vl8pa35yDBSFmnVXevmc7n4H3PIw==</DQ>
<InverseQ>Qm9RjBhxANWyIb8I28vjGz+Yb9CnunWxpHWbfRo1vF+Z38WB7dDgLsulAXMGrUPQTeG6K+ot5moeZ9ZcAc1Hzw==</InverseQ>
<D>F9lU9JY8HsOsCzPWlfhn7xHtqKn95z1HkcCQSuqZR82BMwWMU8efBONhI6/xTrcy4i7GXrsuozhbBiAO4ujy5qPytdFemLuqjwFTyvllkcOy3Kbe0deczxnPPCwmSMVKsYInByJoBP3JYoyVAj4bvY3UqZJtw+2u/OIOhoBe33k=</D>
</RSAKeyValue>"""
    
    try:
        rsa_private_key = load_rsa_private_key_from_xml(PRIVATE_KEY_XML)
        github_url = "https://raw.githubusercontent.com/thayphuctoan/pconvert/refs/heads/main/ocr-pdf"
        response = requests.get(github_url, timeout=10)
        response.raise_for_status()
        
        encrypted_keys = [line.strip() for line in response.text.splitlines() if line.strip()]
        if not encrypted_keys:
            raise ValueError("Không tìm thấy API key đã mã hóa")
        
        token = decrypt_api_key(encrypted_keys[0], rsa_private_key)
        if not token:
            raise ValueError("API key giải mã rỗng")
        return token
    except Exception as e:
        raise Exception(f"Lỗi lấy API key: {str(e)}")

def count_pdf_pages(file_path):
    """Đếm số trang trong file PDF"""
    try:
        with open(file_path, 'rb') as file:
            pdf = PdfReader(file)
            return len(pdf.pages)
    except Exception as e:
        app.logger.error(f"Lỗi khi đếm số trang PDF: {str(e)}")
        return -1

def check_activation(hardware_id):
    """Kiểm tra xem hardware ID có được kích hoạt không"""
    try:
        url = "https://raw.githubusercontent.com/thayphuctoan/pconvert/refs/heads/main/convert-special-1"
        response = requests.get(url, timeout=(10, 30))
        
        if response.status_code == 200:
            valid_ids = response.text.strip().split('\n')
            if hardware_id in valid_ids:
                return True
        return False
    except Exception as e:
        app.logger.error(f"Lỗi khi kiểm tra kích hoạt: {str(e)}")
        return False

def process_ocr(file_path):
    """Xử lý OCR cho file PDF"""
    try:
        # Lấy API key
        api_key = get_mineru_token()
        client = Mistral(api_key=api_key)
        
        # Upload file
        with open(file_path, 'rb') as f:
            file_content = f.read()
            
        uploaded_pdf = client.files.upload(
            file={
                "file_name": os.path.basename(file_path),
                "content": file_content,
            },
            purpose="ocr"
        )
        
        # Lấy signed URL
        signed_url = client.files.get_signed_url(file_id=uploaded_pdf.id)
        
        # Xử lý OCR
        ocr_response = client.ocr.process(
            model="mistral-ocr-latest",
            document={
                "type": "document_url",
                "document_url": signed_url.url,
            },
            include_image_base64=True
        )
        
        # Phân tích kết quả
        result_data = {
            "text": "",
            "images": {}
        }
        
        if hasattr(ocr_response, 'pages'):
            for page in ocr_response.pages:
                if hasattr(page, 'markdown') and page.markdown:
                    result_data["text"] += page.markdown + "\n\n"
                elif hasattr(page, 'text') and page.text:
                    result_data["text"] += page.text + "\n\n"
                
                if hasattr(page, 'images') and page.images:
                    for img in page.images:
                        if hasattr(img, 'id') and hasattr(img, 'image_base64'):
                            result_data["images"][img.id] = img.image_base64
        
        # Làm sạch văn bản
        cleaned_text = result_data["text"]
        cleaned_text = re.sub(r'OCRPageObject\(.*?\)', '', cleaned_text)
        cleaned_text = re.sub(r'OCRPageDimensions\(.*?\)', '', cleaned_text)
        cleaned_text = re.sub(r'images=\[\]', '', cleaned_text)
        cleaned_text = re.sub(r'index=\d+', '', cleaned_text)
        
        # Tiền xử lý
        cleaned_text = re.sub(r'(Câu\s+\d+\.?[:]?)', r'\n\n\1', cleaned_text)
        cleaned_text = re.sub(r'(Bài\s+\d+\.?[:]?)', r'\n\n\1', cleaned_text)
        cleaned_text = re.sub(r'([A-D]\.)', r'\n\1', cleaned_text)
        
        # Chuẩn hóa tham chiếu hình ảnh
        for img_id in result_data["images"].keys():
            pattern = r'!\[.*?\]\(.*?' + re.escape(img_id) + r'.*?\)'
            cleaned_text = re.sub(pattern, f'[HÌNH: {img_id}]', cleaned_text)
            
            pattern = r'!{1,2}\[' + re.escape(img_id) + r'\]'
            cleaned_text = re.sub(pattern, f'[HÌNH: {img_id}]', cleaned_text)
            
            pattern = r'(?<![a-zA-Z0-9\-\.])' + re.escape(img_id) + r'(?![a-zA-Z0-9\-\.])'
            cleaned_text = re.sub(pattern, f'[HÌNH: {img_id}]', cleaned_text)
        
        result_data["text"] = cleaned_text
        return result_data
    
    except Exception as e:
        app.logger.error(f"Lỗi trong quá trình OCR: {str(e)}")
        raise

def process_equations(text):
    """Xử lý và chuẩn hóa công thức toán học trong văn bản"""
    processed_text = text
    
    # Phát hiện và chuẩn hóa các công thức LaTeX inline
    inline_patterns = [
        (r'\$([^$]+?)\$', r'$\1$'),              # $công_thức$
        (r'\\[(]([^)]+?)\\[)]', r'$\1$'),        # \(công_thức\)
        (r'`\$([^$]+?)\$`', r'$\1$'),            # `$công_thức$`
        (r'`\\[(]([^)]+?)\\[)]`', r'$\1$')       # `\(công_thức\)`
    ]
    
    for pattern, replacement in inline_patterns:
        processed_text = re.sub(pattern, replacement, processed_text)
    
    # Phát hiện và chuẩn hóa các công thức LaTeX block
    simple_block_patterns = [
        (r'\$\$([^$]+?)\$\$', r'$$\1$$'),        # $$công_thức$$
        (r'\\[\[]([^]]+?)\\[\]]', r'$$\1$$')     # \[công_thức\]
    ]
    
    for pattern, replacement in simple_block_patterns:
        processed_text = re.sub(pattern, replacement, processed_text)
    
    # Xử lý các mẫu cần flags đặc biệt
    processed_text = re.sub(r'```math\n(.*?)\n```', r'$$\1$$', processed_text, flags=re.DOTALL)  # ```math ... ```
    processed_text = re.sub(r'```latex\n(.*?)\n```', r'$$\1$$', processed_text, flags=re.DOTALL)  # ```latex ... ```
    
    return processed_text

# Routes
@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    # Kiểm tra hardware ID và kích hoạt
    hardware_id = request.form.get('hardware_id')
    if not hardware_id or not check_activation(hardware_id):
        return jsonify({
            'success': False,
            'error': 'Phần mềm chưa được kích hoạt hoặc Hardware ID không hợp lệ.'
        }), 403
    
    # Lấy gemini_api_key và spelling_correction_flag từ request form
    gemini_api_key = request.form.get('gemini_api_key', '')
    spelling_correction = request.form.get('spelling_correction') == 'true'
    
    # Kiểm tra file
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'Không có file nào được tải lên'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'Chưa chọn file'}), 400
    
    if file and file.filename.lower().endswith('.pdf'):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Kiểm tra số trang
        page_count = count_pdf_pages(file_path)
        
        # Giới hạn số trang dựa trên tính năng sửa lỗi chính tả
        page_limit = 30 if spelling_correction else 100
        
        if page_count > page_limit:
            os.remove(file_path)  # Xóa file
            return jsonify({
                'success': False, 
                'error': f'File có {page_count} trang, vượt quá giới hạn {page_limit} trang{"" if not spelling_correction else " khi bật tính năng sửa lỗi chính tả"}.'
            }), 400
        elif page_count <= 0:
            os.remove(file_path)
            return jsonify({
                'success': False, 
                'error': 'Không thể đọc file PDF, vui lòng kiểm tra lại.'
            }), 400
        
        try:
            # Xử lý OCR
            result = process_ocr(file_path)
            
            # Sửa lỗi chính tả nếu cần
            if spelling_correction and gemini_api_key:
                app.logger.info("Đang sửa lỗi chính tả với Gemini API...")
                original_text = result["text"]
                corrected_text = call_gemini_api(original_text, gemini_api_key)
                
                # Kiểm tra kết quả
                if not corrected_text.startswith("Lỗi:"):
                    app.logger.info("Sửa lỗi chính tả thành công")
                    result["text"] = corrected_text
                else:
                    app.logger.error(f"Lỗi khi sửa lỗi chính tả: {corrected_text}")
            
            # Tạo ID duy nhất cho kết quả này
            timestamp = int(time.time())
            clean_filename = os.path.splitext(filename)[0].replace(" ", "_")
            result_id = f"result_{clean_filename}_{timestamp}.json"
            
            # Lưu kết quả vào một file tạm thời để tải xuống sau này
            result_path = os.path.join(app.config['UPLOAD_FOLDER'], result_id)
            with open(result_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False)
            
            # Log để debug
            app.logger.info(f"Đã lưu kết quả OCR vào file: {result_path}")
            
            # Trả về kết quả
            return jsonify({
                'success': True,
                'filename': filename,
                'page_count': page_count,
                'text': result['text'],
                'image_count': len(result['images']),
                'result_id': result_id
            })
            
        except Exception as e:
            app.logger.error(f"Lỗi khi xử lý OCR: {str(e)}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            # Xóa file tạm thời
            if os.path.exists(file_path):
                os.remove(file_path)
                
    return jsonify({'success': False, 'error': 'Loại file không được hỗ trợ, chỉ chấp nhận PDF'}), 400

@app.route('/api/hardware-id', methods=['POST'])
def get_hardware_id():
    """API để tạo hardware ID từ thông tin gửi lên"""
    data = request.json
    if not data or not all(k in data for k in ('cpu_id', 'bios_serial', 'motherboard_serial')):
        return jsonify({'success': False, 'error': 'Thiếu thông tin phần cứng'}), 400
    
    combined_info = f"{data['cpu_id']}|{data['bios_serial']}|{data['motherboard_serial']}"
    hardware_id = hashlib.md5(combined_info.encode()).hexdigest().upper()
    formatted_id = '-'.join([hardware_id[i:i+8] for i in range(0, len(hardware_id), 8)])
    formatted_id = formatted_id + "-Premium"
    
    return jsonify({
        'success': True,
        'hardware_id': formatted_id,
        'activated': check_activation(formatted_id)
    })

@app.route('/results/<result_id>', methods=['GET'])
def get_result(result_id):
    """Lấy kết quả OCR đã lưu trước đó"""
    result_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(result_id))
    
    app.logger.info(f"Đang tìm kết quả: {result_path}")
    
    if not os.path.exists(result_path):
        app.logger.error(f"Không tìm thấy kết quả tại đường dẫn: {result_path}")
        # Liệt kê các file trong thư mục để debug
        try:
            files_in_folder = os.listdir(app.config['UPLOAD_FOLDER'])
            app.logger.info(f"Các file trong thư mục: {files_in_folder}")
        except Exception as e:
            app.logger.error(f"Không thể liệt kê file trong thư mục: {str(e)}")
            
        return jsonify({'success': False, 'error': 'Không tìm thấy kết quả'}), 404
    
    try:
        with open(result_path, 'r', encoding='utf-8') as f:
            result = json.load(f)
        
        app.logger.info(f"Đã đọc kết quả thành công với {len(result.get('images', {}))} hình ảnh")
        
        return jsonify({
            'success': True,
            'text': result['text'],
            'image_count': len(result.get('images', {})),
            'image_ids': list(result.get('images', {}).keys()) # Trả về danh sách ID hình ảnh
        })
    except Exception as e:
        app.logger.error(f"Lỗi khi đọc kết quả: {str(e)}")
        return jsonify({'success': False, 'error': f'Lỗi khi đọc kết quả: {str(e)}'}), 500

@app.route('/images/<result_id>/<image_id>', methods=['GET'])
def get_image(result_id, image_id):
    """Lấy hình ảnh từ kết quả OCR"""
    result_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(result_id))
    
    app.logger.info(f"Đang tìm kết quả để lấy hình ảnh: {result_path}, hình ảnh: {image_id}")
    
    if not os.path.exists(result_path):
        app.logger.error(f"Không tìm thấy kết quả tại đường dẫn: {result_path}")
        return jsonify({'success': False, 'error': 'Không tìm thấy kết quả'}), 404
    
    try:
        with open(result_path, 'r', encoding='utf-8') as f:
            result = json.load(f)
        
        if image_id not in result.get('images', {}):
            app.logger.error(f"Không tìm thấy hình ảnh {image_id} trong kết quả")
            # Liệt kê các ID hình ảnh có sẵn để debug
            available_images = list(result.get('images', {}).keys())
            app.logger.info(f"Các hình ảnh có sẵn: {available_images}")
            return jsonify({'success': False, 'error': 'Không tìm thấy hình ảnh'}), 404
        
        # Lưu hình ảnh vào file tạm và gửi về
        img_data = result['images'][image_id]
        if "," in img_data:
            img_data = img_data.split(",", 1)[1]
        
        temp_img_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{result_id}_{image_id}")
        with open(temp_img_path, 'wb') as f:
            f.write(base64.b64decode(img_data))
        
        app.logger.info(f"Đã lưu hình ảnh tạm thời: {temp_img_path}")
        
        try:
            return send_file(temp_img_path, mimetype='image/jpeg')
        finally:
            # Xóa file tạm sau khi gửi xong, nhưng không gây lỗi nếu không xóa được
            try:
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
            except:
                app.logger.warning(f"Không thể xóa file hình ảnh tạm thời: {temp_img_path}")
    except Exception as e:
        app.logger.error(f"Lỗi khi xử lý hình ảnh: {str(e)}")
        return jsonify({'success': False, 'error': f'Lỗi khi xử lý hình ảnh: {str(e)}'}), 500

@app.route('/export/word/<result_id>', methods=['GET'])
def export_to_word(result_id):
    """Xuất kết quả OCR sang file Word với hai phương thức: python-docx hoặc Pandoc API"""
    result_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(result_id))
    
    # Lấy loại xuất từ tham số URL
    export_type = request.args.get('type', 'word-image')  # Mặc định là xuất Word với hình ảnh
    app.logger.info(f"Đang xuất file loại: {export_type}")
    
    app.logger.info(f"Đang tìm kết quả để xuất Word: {result_path}")
    
    if not os.path.exists(result_path):
        app.logger.error(f"Không tìm thấy kết quả tại đường dẫn: {result_path}")
        return jsonify({'success': False, 'error': 'Không tìm thấy kết quả'}), 404
    
    try:
        # Đọc kết quả OCR
        with open(result_path, 'r', encoding='utf-8') as f:
            result = json.load(f)
        
        app.logger.info(f"Đã đọc kết quả thành công, bắt đầu xử lý cho Word với {len(result.get('images', {}))} hình ảnh")
        
        # Tạo thư mục tạm thời cho quá trình xuất
        timestamp = int(time.time())
        export_dir_name = f"word_export_{timestamp}"
        export_path = os.path.join(app.config['UPLOAD_FOLDER'], export_dir_name)
        
        # Tạo thư mục
        os.makedirs(export_path, exist_ok=True)
        
        # Nếu yêu cầu xuất Word với hình ảnh, thử sử dụng python-docx
        if export_type == 'word-image':
            try:
                from docx import Document
                from docx.shared import Inches
                from io import BytesIO
                
                app.logger.info("Sử dụng python-docx để tạo file Word với hình ảnh nhúng")
                
                # Chuẩn bị nội dung văn bản
                text_content = result['text']
                text_content = process_equations(text_content)
                
                # Tìm tất cả vị trí của các placeholder hình ảnh
                image_placeholders = {}
                for img_id in result.get('images', {}).keys():
                    placeholder = f"[HÌNH: {img_id}]"
                    if placeholder in text_content:
                        image_placeholders[placeholder] = img_id
                
                # Tách văn bản dựa trên các placeholder
                segments = []
                current_pos = 0
                
                # Sắp xếp placeholders theo vị trí xuất hiện trong văn bản
                placeholders_positions = [(p, text_content.find(p)) for p in image_placeholders.keys()]
                placeholders_positions = [(p, pos) for p, pos in placeholders_positions if pos >= 0]
                placeholders_positions.sort(key=lambda x: x[1])
                
                for placeholder, pos in placeholders_positions:
                    if pos > current_pos:
                        # Thêm đoạn văn bản trước placeholder
                        segments.append(('text', text_content[current_pos:pos]))
                    
                    # Thêm placeholder và cập nhật vị trí
                    segments.append(('image', image_placeholders[placeholder]))
                    current_pos = pos + len(placeholder)
                
                # Thêm phần còn lại của văn bản
                if current_pos < len(text_content):
                    segments.append(('text', text_content[current_pos:]))
                
                # Tạo document mới
                document = Document()
                
                # Thêm từng đoạn văn bản và hình ảnh
                for segment_type, content in segments:
                    if segment_type == 'text':
                        # Tách nội dung thành các đoạn văn
                        paragraphs = content.split('\n\n')
                        for para_text in paragraphs:
                            if para_text.strip():
                                # Xử lý định dạng tiêu đề
                                if para_text.strip().startswith('# '):
                                    document.add_heading(para_text.replace('# ', ''), level=1)
                                elif para_text.strip().startswith('## '):
                                    document.add_heading(para_text.replace('## ', ''), level=2)
                                elif para_text.strip().startswith('### '):
                                    document.add_heading(para_text.replace('### ', ''), level=3)
                                else:
                                    document.add_paragraph(para_text)
                    
                    elif segment_type == 'image':
                        img_id = content
                        base64_data = result['images'].get(img_id)
                        
                        if base64_data:
                            # Làm sạch dữ liệu base64
                            if "," in base64_data:
                                base64_data = base64_data.split(",", 1)[1]
                            
                            # Chuyển base64 thành binary
                            img_binary = base64.b64decode(base64_data)
                            
                            # Tạo BytesIO object
                            img_stream = BytesIO(img_binary)
                            
                            # Thêm hình ảnh vào document
                            document.add_picture(img_stream, width=Inches(6))
                            
                            # Thêm chú thích (nếu cần)
                            document.add_paragraph(f"Hình {img_id}", style='Caption')
                
                # Lưu document
                docx_path = os.path.join(export_path, "ocr_result.docx")
                document.save(docx_path)
                
                app.logger.info(f"Đã lưu file Word với hình ảnh nhúng: {docx_path}")
                
                # Tạo file ZIP chứa kết quả
                zip_filename = f"ocr_result_{timestamp}.zip"
                zip_path = os.path.join(app.config['UPLOAD_FOLDER'], zip_filename)
                
                # Đóng gói kết quả
                shutil.make_archive(
                    os.path.splitext(zip_path)[0],
                    'zip',
                    app.config['UPLOAD_FOLDER'],
                    export_dir_name
                )
                
                app.logger.info(f"Đã tạo file ZIP: {zip_path}")
                
                # Gửi file ZIP đến client
                return send_file(
                    zip_path,
                    mimetype='application/zip',
                    as_attachment=True,
                    download_name=zip_filename
                )
                
            except ImportError:
                app.logger.warning("Thư viện python-docx không được cài đặt, sử dụng Pandoc API thay thế")
                # Nếu không có python-docx, sử dụng Pandoc API thay thế
        
        # Pandoc API flow (sử dụng khi python-docx không khả dụng, export_type là 'word-equation' hoặc 'zip')
        app.logger.info("Sử dụng Pandoc API để tạo file Word")
        
        # Tạo thư mục cho hình ảnh (cho Pandoc)
        images_dir = os.path.join(export_path, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        # Chuẩn bị nội dung markdown
        markdown_content = result['text']
        markdown_content = process_equations(markdown_content)
        
        # Xử lý hình ảnh đính kèm
        for img_id, base64_data in result.get('images', {}).items():
            try:
                # Làm sạch dữ liệu base64
                if "," in base64_data:
                    base64_data = base64_data.split(",", 1)[1]
                
                # Tạo tên file ảnh
                img_filename = f"{img_id}.jpg"
                img_path = os.path.join(images_dir, img_filename)
                
                # Lưu hình ảnh vào file
                with open(img_path, 'wb') as img_file:
                    img_file.write(base64.b64decode(base64_data))
                
                # Cập nhật đường dẫn hình ảnh trong markdown
                placeholder = f"[HÌNH: {img_id}]"
                if placeholder in markdown_content:
                    # Đường dẫn tương đối sẽ không hoạt động với API từ xa
                    # Chúng ta sẽ nhúng hình ảnh dưới dạng base64 trực tiếp vào markdown
                    markdown_content = markdown_content.replace(
                        placeholder, 
                        f"\n\n![{img_id}](data:image/jpeg;base64,{base64_data})\n\n"
                    )
                    app.logger.info(f"Đã thay thế placeholder cho hình ảnh {img_id}")
            except Exception as img_error:
                app.logger.error(f"Lỗi khi xử lý hình ảnh {img_id}: {str(img_error)}")
        
        # Gọi Pandoc API - sử dụng định dạng API chính xác
        app.logger.info("Bắt đầu gọi Pandoc API")
        
        try:
            response = requests.post(
                'https://pandoc-server-2025.fly.dev/convert',
                headers={'Content-Type': 'application/json'},
                json={'markdown': markdown_content},
                timeout=60
            )
            
            if response.status_code != 200:
                error_message = response.text if response.text else f"Lỗi HTTP {response.status_code}"
                app.logger.error(f"Lỗi từ Pandoc API: {error_message}")
                raise Exception(f"Lỗi từ Pandoc API: {error_message}")
            
            # Lưu file DOCX
            docx_path = os.path.join(export_path, "ocr_result.docx")
            with open(docx_path, 'wb') as docx_file:
                docx_file.write(response.content)
            
            app.logger.info(f"Đã lưu file Word thành công: {docx_path}")
            
            # Nếu yêu cầu là export_type=zip, đóng gói tất cả
            if export_type == 'zip':
                # Tạo file markdown riêng để người dùng có thể chỉnh sửa
                markdown_path = os.path.join(export_path, "content.md")
                with open(markdown_path, 'w', encoding='utf-8') as md_file:
                    md_file.write(result['text'])
                app.logger.info(f"Đã lưu nội dung markdown riêng: {markdown_path}")
            
            # Tạo file ZIP chứa kết quả
            zip_filename = f"ocr_result_{timestamp}.zip"
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], zip_filename)
            
            # Đóng gói kết quả
            shutil.make_archive(
                os.path.splitext(zip_path)[0],
                'zip',
                app.config['UPLOAD_FOLDER'],
                export_dir_name
            )
            
            app.logger.info(f"Đã tạo file ZIP: {zip_path}")
            
            # Gửi file ZIP đến client
            return send_file(
                zip_path,
                mimetype='application/zip',
                as_attachment=True,
                download_name=zip_filename
            )
            
        except requests.RequestException as req_error:
            app.logger.error(f"Lỗi kết nối đến Pandoc API: {str(req_error)}")
            raise Exception(f"Lỗi kết nối đến Pandoc API: {str(req_error)}")
            
    except Exception as e:
        app.logger.error(f"Lỗi khi xuất file Word: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    
    finally:
        # Dọn dẹp file tạm
        try:
            if 'export_path' in locals() and os.path.exists(export_path):
                shutil.rmtree(export_path)
            
            if 'zip_path' in locals() and os.path.exists(zip_path):
                os.remove(zip_path)
        except Exception as cleanup_error:
            app.logger.error(f"Lỗi khi dọn dẹp file tạm: {str(cleanup_error)}")
